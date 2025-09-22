from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# Membuat dokumen Word
doc = Document()

# Judul
doc.add_heading('Algoritma Pencarian Linear: Studi Kasus Kartu Parkir Hilang', level=1)

# Deskripsi Kasus
doc.add_heading('1. Deskripsi Kasus', level=2)
doc.add_paragraph(
    "Di sebuah gedung perkantoran, petugas parkir menyimpan struk parkir harian yang sudah tidak terpakai dalam sebuah kotak. Suatu hari, seorang pengunjung kehilangan kartu parkirnya dan meminta bantuan untuk mencarinya. Karena struk disimpan secara acak dan tidak berurutan, satu-satunya cara untuk menemukan kartu tersebut adalah dengan memeriksa setiap struk satu per satu. Pendekatan ini merupakan contoh nyata dari algoritma pencarian linear (linear search)."
)

# Pseudocode
doc.add_heading('2. Pseudocode', level=2)
doc.add_paragraph(
    "LinearSearch(array, target):\n"
    "    for i from 0 to length(array) - 1:\n"
    "        if array[i] == target:\n"
    "            return i\n"
    "    return -1"
)

# Kode Python
doc.add_heading('3. Implementasi Python', level=2)
doc.add_paragraph(
    "def linear_search(array, target):\n"
    "    for i in range(len(array)):\n"
    "        if array[i] == target:\n"
    "            return i\n"
    "    return -1\n\n"
    "# Contoh penggunaan\n"
    "tumpukan_struk = [\"A-102\", \"B-215\", \"C-330\", \"A-145\", \"D-501\", \"B-215\", \"E-777\"]\n"
    "target_struk = \"D-501\"\n"
    "hasil = linear_search(tumpukan_struk, target_struk)\n\n"
    "if hasil != -1:\n"
    "    print(f'Struk parkir ditemukan pada indeks ke-{hasil}')\n"
    "else:\n"
    "    print('Struk parkir tidak ditemukan')"
)

# Analisis Kompleksitas
doc.add_heading('4. Analisis Kompleksitas Waktu (Big O)', level=2)
doc.add_paragraph(
    "- Kasus Terbaik (O(1)): Jika struk yang dicari berada di posisi pertama dalam daftar, maka pencarian selesai dalam satu langkah.\n"
    "- Kasus Terburuk (O(n)): Jika struk berada di posisi terakhir atau tidak ada dalam daftar, maka seluruh elemen harus diperiksa.\n"
    "- Kasus Rata-rata (O(n/2)): Dalam kondisi rata-rata, struk ditemukan di tengah daftar. Namun dalam notasi Big O, ini tetap disederhanakan menjadi O(n)."
)

# Penjelasan Esai
doc.add_heading('5. Penjelasan dalam Bentuk Esai', level=2)
doc.add_paragraph(
    "Algoritma pencarian linear adalah metode pencarian yang paling sederhana, di mana setiap elemen dalam daftar diperiksa satu per satu hingga ditemukan kecocokan atau hingga seluruh daftar selesai diperiksa.\n\n"
    "Dalam studi kasus ini, petugas parkir mencoba menemukan kartu parkir yang hilang di antara tumpukan struk harian. Karena struk disimpan secara acak dan tidak berurutan, pencarian harus dilakukan secara linear.\n\n"
    "Jika kartu yang dicari berada di awal tumpukan, pencarian akan sangat cepat (O(1)). Namun jika berada di akhir atau tidak ada sama sekali, pencarian menjadi lambat karena harus memeriksa semua struk (O(n)). Dalam kondisi rata-rata, pencarian akan menemukan kartu di tengah-tengah daftar, yang secara matematis memerlukan n/2 langkah, tetapi tetap dikategorikan sebagai O(n) dalam notasi Big O.\n\n"
    "Metode ini cocok untuk daftar kecil atau data yang tidak terstruktur, namun kurang efisien untuk daftar besar. Studi kasus ini menunjukkan bagaimana algoritma sederhana seperti linear search tetap relevan dalam situasi nyata yang tidak terstruktur."
)

# Simpan dokumen
output_path = "/mnt/data/Esai_Linear_Search_Kartu_Parkir.docx"
doc.save(output_path)
print(f"Dokumen berhasil dibuat: {output_path}")


